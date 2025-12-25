from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import io
import os
import pypandoc
import requests
from PIL import Image
import tempfile
import lxml.etree as ET

app = Flask(__name__)
CORS(app)  # Enable CORS for frontend requests

# Direct download URLs for stamp and signature
PAID_STAMP_URL = "https://drive.google.com/uc?export=download&id=1W9PL0DtP0TUk7IcGiMD_ZuLddtQ8gjNo"
SIGNATURE_URL = "https://drive.google.com/uc?export=download&id=1b6Dcg4spQmvLUMd4neBtLNfdr5l7QtPJ"

# === UTILITY FUNCTIONS ===

def format_currency(amount):
    if amount == 0:
        return ""
    elif amount == int(amount):
        return f"Rp {int(amount):,}"
    else:
        return f"Rp {amount:,.2f}"

def set_cell_border(cell, side, color="FFFFFF", sz=4):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    side_mapping = {
        'top': 'top', 'bottom': 'bottom', 'left': 'left', 'right': 'right'
    }
    border_name = side_mapping.get(side.lower())
    if border_name:
        border = parse_xml(f'<w:{border_name} {nsdecls("w")} w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>')
        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        if tcBorders is None:
            tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
            tcPr.append(tcBorders)
        tcBorders.append(border)

def set_white_borders(cell, sz=4):
    for border in ['top', 'bottom', 'left', 'right']:
        set_cell_border(cell, border, color="FFFFFF", sz=sz)

def set_cell_font(cell, font_name="Courier New", font_size=10):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def apply_cell_style(cell, bg_color="#ddefd5"):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}" />')
    cell._tc.get_or_add_tcPr().append(shading_elm)
    set_white_borders(cell, sz=6)
    set_cell_font(cell)

def replace_placeholders(doc, replacements):
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replacements.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)
    return doc

def update_items_table(doc, items):
    items_table = doc.tables[0]
    for i in range(len(items_table.rows)):
        for cell in items_table.rows[i].cells:
            set_white_borders(cell, sz=6)
    while len(items_table.rows) > 2:
        items_table._tbl.remove(items_table.rows[2]._tr)
    placeholder_row = items_table.rows[1]
    for item in items:
        row = items_table.add_row()
        row.cells[0].text = item['description']
        row.cells[1].text = format_currency(item['unit_price'])
        quantity = item['quantity']
        if quantity == int(quantity):
            row.cells[2].text = str(int(quantity))
        else:
            row.cells[2].text = str(quantity)
        row.cells[3].text = format_currency(item['total'])
        for i, cell in enumerate(row.cells):
            apply_cell_style(cell)
            alignments = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT, 
                         WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT]
            for paragraph in cell.paragraphs:
                paragraph.alignment = alignments[i]
    items_table._tbl.remove(placeholder_row._tr)
    return doc

def style_financial_table(doc, apply_late_fee):
    financial_table = doc.tables[1]
    for row in financial_table.rows:
        for cell in row.cells:
            set_white_borders(cell)
            set_cell_font(cell)
        for paragraph in row.cells[1].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if apply_late_fee:
        late_fee_cell = financial_table.rows[3].cells[0]
        if "LATE FEE" in late_fee_cell.text:
            original_text = late_fee_cell.text
            late_fee_cell.text = ""
            paragraph = late_fee_cell.paragraphs[0]
            run = paragraph.add_run(original_text)
            run.font.color.rgb = RGBColor.from_string('d95132')
            run.font.name = "Courier New"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "Courier New")

def fetch_image(url):
    try:
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
        }
        response = requests.get(url, headers=headers, stream=True, allow_redirects=True)
        if response.status_code != 200:
            raise Exception(f"Failed to fetch image. Status: {response.status_code}")
        image_data = io.BytesIO(response.content)
        img = Image.open(image_data)
        img.verify()
        image_data.seek(0)
        return image_data
    except Exception as e:
        raise Exception(f"Error fetching image: {str(e)}")

def add_paid_stamp_and_signature(doc):
    try:
        stamp_data = fetch_image(PAID_STAMP_URL)
        signature_data = fetch_image(SIGNATURE_URL)

        stamp_tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
        signature_tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.png')

        stamp_img = Image.open(stamp_data)
        stamp_img.save(stamp_tmp.name, format="PNG")
        stamp_tmp.close()

        signature_img = Image.open(signature_data)
        signature_img.save(signature_tmp.name, format="PNG")
        signature_tmp.close()

        # Add stamp
        stamp_paragraph = doc.add_paragraph()
        stamp_run = stamp_paragraph.add_run()
        stamp_run.add_picture(stamp_tmp.name, width=Inches(2.17), height=Inches(2.17))
        
        stamp_run_element = stamp_run._r
        stamp_drawing = stamp_run_element.xpath('.//w:drawing')[0]
        graphic_elements = stamp_drawing.xpath('.//a:graphic', namespaces={'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})
        graphic_xml = ET.tostring(graphic_elements[0], encoding='unicode').replace('\n', '')

        stamp_horizontal = 5.09 * 914400
        stamp_vertical = 6.64 * 914400

        stamp_drawing.getparent().replace(stamp_drawing, parse_xml(f"""
            <w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing">
                <wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="251" behindDoc="0" locked="0" layoutInCell="1" allowOverlap="1">
                    <wp:simplePos x="0" y="0"/>
                    <wp:positionH relativeFrom="page">
                        <wp:posOffset>{int(stamp_horizontal)}</wp:posOffset>
                    </wp:positionH>
                    <wp:positionV relativeFrom="page">
                        <wp:posOffset>{int(stamp_vertical)}</wp:posOffset>
                    </wp:positionV>
                    <wp:extent cx="{int(2.17 * 914400)}" cy="{int(2.17 * 914400)}"/>
                    <wp:effectExtent l="0" t="0" r="0" b="0"/>
                    <wp:wrapTopAndBottom/>
                    <wp:docPr id="1" name="Picture 1"/>
                    <wp:cNvGraphicFramePr/>
                    {graphic_xml}
                </wp:anchor>
            </w:drawing>
        """))

        # Add signature
        signature_paragraph = doc.add_paragraph()
        signature_run = signature_paragraph.add_run()
        signature_run.add_picture(signature_tmp.name, width=Inches(1.92), height=Inches(1.92))

        signature_run_element = signature_run._r
        signature_drawing = signature_run_element.xpath('.//w:drawing')[0]
        graphic_elements = signature_drawing.xpath('.//a:graphic', namespaces={'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})
        graphic_xml = ET.tostring(graphic_elements[0], encoding='unicode').replace('\n', '')

        signature_horizontal = 5.64 * 914400
        signature_vertical = 8.11 * 914400

        signature_drawing.getparent().replace(signature_drawing, parse_xml(f"""
            <w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing">
                <wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="252" behindDoc="0" locked="0" layoutInCell="1" allowOverlap="1">
                    <wp:simplePos x="0" y="0"/>
                    <wp:positionH relativeFrom="page">
                        <wp:posOffset>{int(signature_horizontal)}</wp:posOffset>
                    </wp:positionH>
                    <wp:positionV relativeFrom="page">
                        <wp:posOffset>{int(signature_vertical)}</wp:posOffset>
                    </wp:positionV>
                    <wp:extent cx="{int(1.92 * 914400)}" cy="{int(1.92 * 914400)}"/>
                    <wp:effectExtent l="0" t="0" r="0" b="0"/>
                    <wp:wrapTopAndBottom/>
                    <wp:docPr id="2" name="Picture 2"/>
                    <wp:cNvGraphicFramePr/>
                    {graphic_xml}
                </wp:anchor>
            </w:drawing>
        """))

        os.remove(stamp_tmp.name)
        os.remove(signature_tmp.name)

        return doc
    except Exception as e:
        if 'stamp_tmp' in locals() and os.path.exists(stamp_tmp.name):
            os.remove(stamp_tmp.name)
        if 'signature_tmp' in locals() and os.path.exists(signature_tmp.name):
            os.remove(signature_tmp.name)
        raise Exception(f"Failed to add stamp and signature: {str(e)}")

# === API ENDPOINTS ===

@app.route('/')
def home():
    return jsonify({
        "message": "Invoice Generator API",
        "status": "running",
        "endpoints": {
            "/health": "Check API health",
            "/generate-invoice": "POST - Generate invoice (DOCX and PDF)"
        }
    })

@app.route('/health')
def health():
    return jsonify({"status": "healthy"})

@app.route('/generate-invoice', methods=['POST'])
def generate_invoice():
    try:
        data = request.json
        
        # Extract data from request
        client_info = data.get('client_info', {})
        invoice_details = data.get('invoice_details', {})
        items = data.get('items', [])
        financials = data.get('financials', {})
        apply_late_fee = data.get('apply_late_fee', False)
        mark_as_paid = data.get('mark_as_paid', False)
        output_format = data.get('format', 'docx')  # 'docx' or 'pdf'
        
        # Load template
        doc = Document('Invoice_Template_MarketixLab.docx')
        
        # Prepare replacements
        replacements = {**client_info, **invoice_details, **financials}
        if apply_late_fee:
            replacements['{{LATE FEE:}}'] = 'LATE FEE'
        else:
            replacements['{{LATE FEE:}}'] = ''
            replacements['[latefee]'] = ''
        
        # Apply replacements
        doc = replace_placeholders(doc, replacements)
        doc = update_items_table(doc, items)
        style_financial_table(doc, apply_late_fee)
        
        # Add paid stamp if needed
        if mark_as_paid:
            doc = add_paid_stamp_and_signature(doc)
        
        # Set font for all paragraphs
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.font.name = "Courier New"
                run._element.rPr.rFonts.set(qn('w:eastAsia'), "Courier New")
        
        # Generate output based on format
        if output_format == 'pdf':
            # Save as temporary DOCX first
            temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
            doc.save(temp_docx.name)
            temp_docx.close()
            
            # Convert to PDF
            temp_pdf = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
            temp_pdf.close()
            pypandoc.convert_file(temp_docx.name, 'pdf', outputfile=temp_pdf.name)
            
            # Clean up DOCX
            os.remove(temp_docx.name)
            
            # Return PDF
            return send_file(
                temp_pdf.name,
                mimetype='application/pdf',
                as_attachment=True,
                download_name=f"Invoice_{invoice_details.get('{{invoice_number}}', 'unknown')}.pdf"
            )
        else:
            # Return DOCX
            docx_output = io.BytesIO()
            doc.save(docx_output)
            docx_output.seek(0)
            
            return send_file(
                docx_output,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=f"Invoice_{invoice_details.get('{{invoice_number}}', 'unknown')}.docx"
            )
            
    except Exception as e:
        return jsonify({"error": str(e)}), 500

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 10000))
    app.run(host='0.0.0.0', port=port, debug=False)